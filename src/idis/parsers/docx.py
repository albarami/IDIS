"""DOCX parser — deterministic text extraction with paragraph locators.

Extracts text from DOCX files paragraph by paragraph and from tables,
producing SpanDraft objects with stable locators.

Requirements:
- Deterministic: same bytes in → same ordered spans out
- Fail-closed: malformed DOCX returns structured errors
- Stable locators: paragraphs indexed in document order
"""

from __future__ import annotations

import hashlib
import io
from typing import TYPE_CHECKING

from idis.parsers.base import (
    ParseError,
    ParseErrorCode,
    ParseLimits,
    ParseResult,
    SpanDraft,
)

if TYPE_CHECKING:
    pass

from docx import Document
from docx.opc.exceptions import PackageNotFoundError


def _compute_content_hash(text: str) -> str:
    """Compute SHA-256 hash of text content."""
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _normalize_text(text: str) -> str:
    """Normalize text for consistent output.

    - Convert \\r\\n and \\r to \\n
    - Strip trailing whitespace from each line
    - Strip leading/trailing whitespace from result
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    return "\n".join(lines).strip()


def parse_docx(
    data: bytes,
    limits: ParseLimits | None = None,
) -> ParseResult:
    """Parse DOCX bytes and extract text spans with paragraph/table locators.

    Args:
        data: Raw DOCX file bytes.
        limits: Optional parsing limits (defaults to ParseLimits()).

    Returns:
        ParseResult with success=True and spans if extraction succeeded,
        or success=False with structured errors if parsing failed.

    Behavior:
        - Paragraphs are processed in document order (0-indexed).
        - Tables are processed in document order after paragraphs.
        - Each paragraph becomes a SpanDraft with locator {paragraph: i}.
        - Each table cell becomes a SpanDraft with locator {table: t, row: r, col: c}.
        - Empty paragraphs/cells are skipped.
        - Malformed DOCX files fail with CORRUPTED_FILE error.
    """
    if limits is None:
        limits = ParseLimits()

    if len(data) > limits.max_bytes:
        return ParseResult(
            doc_type="DOCX",
            success=False,
            errors=[
                ParseError(
                    code=ParseErrorCode.MAX_SIZE_EXCEEDED,
                    message=f"File size {len(data)} bytes exceeds limit {limits.max_bytes}",
                    details={"size": len(data), "limit": limits.max_bytes},
                )
            ],
        )

    try:
        doc = Document(io.BytesIO(data))
    except PackageNotFoundError as e:
        return ParseResult(
            doc_type="DOCX",
            success=False,
            errors=[
                ParseError(
                    code=ParseErrorCode.CORRUPTED_FILE,
                    message="Invalid DOCX file: not a valid Office Open XML package",
                    details={"error": str(e)},
                )
            ],
        )
    except Exception as e:
        return ParseResult(
            doc_type="DOCX",
            success=False,
            errors=[
                ParseError(
                    code=ParseErrorCode.CORRUPTED_FILE,
                    message="Failed to read DOCX file",
                    details={"error": str(e), "type": type(e).__name__},
                )
            ],
        )

    spans: list[SpanDraft] = []
    warnings: list[str] = []
    total_text_length = 0
    paragraph_count = 0
    table_count = 0

    for para_idx, paragraph in enumerate(doc.paragraphs):
        text = _normalize_text(paragraph.text)
        if not text:
            continue

        total_text_length += len(text)
        paragraph_count += 1

        spans.append(
            SpanDraft(
                span_type="PARAGRAPH",
                locator={"paragraph": para_idx},
                text_excerpt=text,
                content_hash=_compute_content_hash(text),
            )
        )

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                text = _normalize_text(cell.text)
                if not text:
                    continue

                total_text_length += len(text)
                table_count += 1

                spans.append(
                    SpanDraft(
                        span_type="CELL",
                        locator={
                            "table": table_idx,
                            "row": row_idx,
                            "col": col_idx,
                        },
                        text_excerpt=text,
                        content_hash=_compute_content_hash(text),
                    )
                )

    if total_text_length == 0:
        return ParseResult(
            doc_type="DOCX",
            success=False,
            errors=[
                ParseError(
                    code=ParseErrorCode.NO_TEXT_EXTRACTED,
                    message="No extractable text found in DOCX",
                    details={
                        "paragraph_count": len(doc.paragraphs),
                        "table_count": len(doc.tables),
                    },
                )
            ],
        )

    return ParseResult(
        doc_type="DOCX",
        success=True,
        spans=spans,
        metadata={
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "span_count": len(spans),
            "total_text_length": total_text_length,
        },
        warnings=warnings,
    )
