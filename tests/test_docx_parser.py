"""Tests for DOCX parser — deterministic extraction with stable locators.

Tests verify:
- Successful parsing with correct locators
- Deterministic output (same input → same spans)
- Fail-closed behavior for corrupted files
- Table extraction with row/col locators
"""

from __future__ import annotations

import io

from docx import Document

from idis.parsers.base import ParseErrorCode, ParseLimits
from idis.parsers.docx import parse_docx


def create_test_docx(paragraphs: list[str], tables: list[list[list[str]]] | None = None) -> bytes:
    """Create a DOCX file in memory with specified paragraphs and tables.

    Args:
        paragraphs: List of paragraph texts.
        tables: Optional list of tables, each as [[row1_cells], [row2_cells], ...].

    Returns:
        DOCX file as bytes.
    """
    doc = Document()

    for para_text in paragraphs:
        doc.add_paragraph(para_text)

    if tables:
        for table_data in tables:
            if not table_data:
                continue
            num_rows = len(table_data)
            num_cols = max(len(row) for row in table_data) if table_data else 0
            if num_cols == 0:
                continue

            table = doc.add_table(rows=num_rows, cols=num_cols)
            for row_idx, row_data in enumerate(table_data):
                for col_idx, cell_text in enumerate(row_data):
                    table.rows[row_idx].cells[col_idx].text = cell_text

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class TestDOCXParserSuccess:
    """Test successful DOCX parsing."""

    def test_parse_simple_docx(self) -> None:
        """Parse DOCX with paragraphs returns success."""
        paragraphs = ["First paragraph.", "Second paragraph."]
        docx_bytes = create_test_docx(paragraphs)

        result = parse_docx(docx_bytes)

        assert result.success is True
        assert result.doc_type == "DOCX"
        assert len(result.errors) == 0
        assert len(result.spans) == 2

    def test_parse_docx_with_tables(self) -> None:
        """Parse DOCX with tables extracts cells."""
        paragraphs = ["Header text"]
        tables = [[["A1", "B1"], ["A2", "B2"]]]
        docx_bytes = create_test_docx(paragraphs, tables)

        result = parse_docx(docx_bytes)

        assert result.success is True
        paragraph_spans = [s for s in result.spans if s.span_type == "PARAGRAPH"]
        cell_spans = [s for s in result.spans if s.span_type == "CELL"]
        assert len(paragraph_spans) == 1
        assert len(cell_spans) == 4

    def test_metadata_populated(self) -> None:
        """Metadata contains paragraph and span counts."""
        paragraphs = ["Para 1", "Para 2", "Para 3"]
        docx_bytes = create_test_docx(paragraphs)

        result = parse_docx(docx_bytes)

        assert result.metadata["paragraph_count"] == 3
        assert result.metadata["span_count"] == 3
        assert result.metadata["total_text_length"] > 0


class TestDOCXParserLocators:
    """Test locator correctness for DOCX spans."""

    def test_paragraph_locators_zero_indexed(self) -> None:
        """Paragraph locators are 0-indexed in document order."""
        paragraphs = ["First", "Second", "Third"]
        docx_bytes = create_test_docx(paragraphs)

        result = parse_docx(docx_bytes)

        assert result.success is True
        locators = [s.locator for s in result.spans if s.span_type == "PARAGRAPH"]
        assert locators[0]["paragraph"] == 0
        assert locators[1]["paragraph"] == 1
        assert locators[2]["paragraph"] == 2

    def test_table_cell_locators(self) -> None:
        """Table cell locators include table, row, col."""
        tables = [[["R0C0", "R0C1"], ["R1C0", "R1C1"]]]
        docx_bytes = create_test_docx([], tables)

        result = parse_docx(docx_bytes)

        assert result.success is True
        cell_spans = [s for s in result.spans if s.span_type == "CELL"]
        assert len(cell_spans) == 4

        locator_0 = cell_spans[0].locator
        assert "table" in locator_0
        assert "row" in locator_0
        assert "col" in locator_0
        assert locator_0["table"] == 0
        assert locator_0["row"] == 0
        assert locator_0["col"] == 0

    def test_content_hash_populated(self) -> None:
        """Each span has a content hash."""
        paragraphs = ["Test content"]
        docx_bytes = create_test_docx(paragraphs)

        result = parse_docx(docx_bytes)

        assert result.success is True
        for span in result.spans:
            assert span.content_hash is not None
            assert len(span.content_hash) == 64  # SHA-256 hex


class TestDOCXParserDeterminism:
    """Test deterministic output."""

    def test_same_input_same_output(self) -> None:
        """Parsing same bytes twice produces identical spans."""
        paragraphs = ["Revenue: $10M", "Growth: 150%"]
        docx_bytes = create_test_docx(paragraphs)

        result1 = parse_docx(docx_bytes)
        result2 = parse_docx(docx_bytes)

        assert result1.success is True
        assert result2.success is True
        assert len(result1.spans) == len(result2.spans)

        for span1, span2 in zip(result1.spans, result2.spans, strict=True):
            assert span1.locator == span2.locator
            assert span1.text_excerpt == span2.text_excerpt
            assert span1.content_hash == span2.content_hash
            assert span1.span_type == span2.span_type


class TestDOCXParserFailClosed:
    """Test fail-closed behavior for invalid inputs."""

    def test_corrupted_bytes(self) -> None:
        """Corrupted data returns error, no exception."""
        corrupted = b"PK\x03\x04not_a_real_docx_file"

        result = parse_docx(corrupted)

        assert result.success is False
        assert result.doc_type == "DOCX"
        assert len(result.errors) > 0
        assert result.errors[0].code == ParseErrorCode.CORRUPTED_FILE

    def test_random_bytes(self) -> None:
        """Random bytes return corrupted file error."""
        random_data = b"\x00\x01\x02\x03\x04\x05\x06\x07\x08\x09"

        result = parse_docx(random_data)

        assert result.success is False
        assert result.errors[0].code == ParseErrorCode.CORRUPTED_FILE

    def test_empty_docx_no_text(self) -> None:
        """DOCX with no text returns NO_TEXT_EXTRACTED."""
        docx_bytes = create_test_docx([])

        result = parse_docx(docx_bytes)

        assert result.success is False
        assert result.errors[0].code == ParseErrorCode.NO_TEXT_EXTRACTED


class TestDOCXParserLimits:
    """Test limit enforcement."""

    def test_max_size_exceeded(self) -> None:
        """File exceeding max_bytes returns error."""
        paragraphs = ["Test paragraph"]
        docx_bytes = create_test_docx(paragraphs)

        limits = ParseLimits(max_bytes=100)  # Very small limit
        result = parse_docx(docx_bytes, limits=limits)

        assert result.success is False
        assert result.errors[0].code == ParseErrorCode.MAX_SIZE_EXCEEDED


class TestDOCXParserEdgeCases:
    """Test edge cases."""

    def test_empty_paragraphs_skipped(self) -> None:
        """Empty paragraphs are not included in spans."""
        paragraphs = ["First", "", "Third", ""]
        docx_bytes = create_test_docx(paragraphs)

        result = parse_docx(docx_bytes)

        assert result.success is True
        assert len(result.spans) == 2
        texts = [s.text_excerpt for s in result.spans]
        assert "First" in texts
        assert "Third" in texts

    def test_whitespace_only_paragraphs_skipped(self) -> None:
        """Whitespace-only paragraphs are skipped."""
        paragraphs = ["Content", "   ", "\t\n", "More content"]
        docx_bytes = create_test_docx(paragraphs)

        result = parse_docx(docx_bytes)

        assert result.success is True
        assert len(result.spans) == 2

    def test_unicode_content(self) -> None:
        """Unicode content is preserved."""
        paragraphs = ["Hello 世界", "مرحبا بالعالم", "🎉 Celebration"]
        docx_bytes = create_test_docx(paragraphs)

        result = parse_docx(docx_bytes)

        assert result.success is True
        assert len(result.spans) == 3
        texts = [s.text_excerpt for s in result.spans]
        assert "Hello 世界" in texts
        assert "مرحبا بالعالم" in texts
        assert "🎉 Celebration" in texts
